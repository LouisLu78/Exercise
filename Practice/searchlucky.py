# -*- coding:utf8 -*-
import bs4, webbrowser, sys, requests

# print('searching in baidu...')
# res=requests.get('http://www.baidu.com/s?wd=' + ' '.join(sys.argv[1:]))
# res.raise_for_status()
# # print(res.text[:1000])
# soup=bs4.BeautifulSoup(res.text)
# print(soup)
# linkElems = soup.select('.b a')
# numOpen = min(5, len(linkElems))
# for i in range(numOpen):
#     webbrowser.open('http://www.baidu.com' + linkElems[i].get('href'))

# from selenium import webdriver
# browser=webdriver.Edge()
# browser.get('www.python.org')

# import docx
# mydoc123=docx.Document("D:\\Arbeiten\\myfiles\\Arbeiten\\学术评论之一百二十三.docx")
# print('Mydoc123 contains {} paragraghs'.format(len(mydoc123.paragraphs)))
# for i in range(len(mydoc123.paragraphs)):
#     print('This is the paragraph {}...'.format(i+1))
#     print(mydoc123.paragraphs[i].text, end='\n')
# print(mydoc123.paragraphs[1].text)
# print('Paragragh 1 contains {} runs'.format(len(mydoc123.paragraphs[1].runs)))
# print(mydoc123.paragraphs[1].runs[1].text)
#
# def retrive_text(file):
#     doctext=docx.Document(file)
#     fulltext=[]
#     for para in doctext.paragraphs:
#         fulltext.append(para.text)
#     return '\n'.join(fulltext)
# doc123= retrive_text("D:\\Arbeiten\\myfiles\\Arbeiten\\学术评论之一百二十三.docx")
# print(doc123)
# print(len(str(doc123)))
import docx
with open('downfile.txt','r') as f:
    cont=f.read()
newdoc=docx.Document()
newdoc.add_paragraph(cont)
newdoc.save('downfile.docx')

# def cal(n):
#     result=1
#     for i in range(1,n):
#         result*=i
#     return result
# import time
# start=time.time()
# product=cal(100000)
# end=time.time()
# print('The result has {} digits.'.format(len(str(product))))
# print('It takes {} seconds to finish the process.'.format(end-start))

import datetime
dt=datetime.datetime.today()
now=datetime.datetime.now()
print(dt)
print(now)

import subprocess
subprocess.Popen(['C:\\Windows\\notepad.exe', 'C:\\hello.txt'])


